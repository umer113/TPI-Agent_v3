import asyncio
import os
import shutil
import subprocess
from datetime import datetime
from dotenv import load_dotenv
import pandas as pd
import streamlit as st
from openai import AsyncOpenAI
from groq import Groq
import json
from bs4 import BeautifulSoup
from urllib.parse import urljoin
import csv
import io
import requests
import hashlib
import sys
from docx import Document
from io import BytesIO
import random
import tiktoken
import shutil

load_dotenv()

FIRST_PAGE_URL = {
    "DVA Minister": "https://minister.dva.gov.au/minister-media-releases?page=1",
    "DVA Veteran Affairs": "https://www.dva.gov.au/about/news/vetaffairs",
    "DVA Repatriation Commission": "https://www.dva.gov.au/about/overview/repatriation-commission/gwen-cherne-veteran-family-advocate-commissioner/veteran-family-advocate-commissioner-gwen-cherne",
    "DVA Website About": "https://www.dva.gov.au/about/our-work-response-royal-commission-defence-and-veteran-suicide",
    "DVA Website Home": "https://clik.dva.gov.au/",
    "DVA Website Latest News": "https://www.dva.gov.au/about/news/latest-news",
    "articles":"https://www.awm.gov.au/articles",
    "RMA":"http://www.rma.gov.au/",
    "X AWM": "https://x.com/awmemorial?lang=en",
    "X DVA": "https://x.com/DVAAus",
    "Instagram DVA": "https://www.instagram.com/dvaausgov/",
    # "Instagram AWM" : 
}

TOP5_SELECTORS = {
    # Australian War Memorial: each card is in a <div class="article--card">
    "articles":      "div.article--card a",

    # (omit or set to None if you want RMA to always do a full scrape)
    "RMA":                          None,

    # DVA Minister releases: each title block has this class
    "DVA Minister":                 "div.media_release_listing--content-title a",

    # DVA Veteran Affairs: listings are in <div class="col-md-6"> with <a class="card">
    "DVA Veteran Affairs":          "div.col-md-6 a.card",

    "X AWM" : None,

    "X DVA" : None,

    # "Instagram AWM": None,

    "Instagram DVA":None,


    # All the other DVA pages share the same “views” layout
    "DVA Repatriation Commission":  "div.views-row h2 a",
    "DVA Website About":            "div.views-row h2 a",
    "DVA Website Home":             "div.views-row h2 a",

    # DVA “Latest News” uses simple cards
    "DVA Website Latest News":      "a.card[href]",

    
}

# Fallback if a key isn’t found:
DEFAULT_SELECTOR = "a.card"

API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    st.error("Please set the OPENAI_API_KEY in your .env file.")
    st.stop()

OPENAI_API_KEY = API_KEY
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

CHAT_DIR = "chat_history"
DATA_DIR = "data"
os.makedirs(CHAT_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)

# ——— Chat Persistence ———
def _clean_filename(s: str) -> str:
    return "".join(c if c.isalnum() or c in "-_ " else "_" for c in s).strip().replace(" ", "_")

def save_chat(history, chat_id=None, title=None):
    if title:
        new_id = _clean_filename(title)[:50] or datetime.now().strftime("%Y%m%d_%H%M%S")
    elif chat_id:
        new_id = chat_id
    else:
        new_id = datetime.now().strftime("%Y%m%d_%H%M%S")

    new_path = os.path.join(CHAT_DIR, f"{new_id}.json")
    if chat_id and chat_id != new_id:
        old_path = os.path.join(CHAT_DIR, f"{chat_id}.json")
        if os.path.exists(old_path):
            os.remove(old_path)

    with open(new_path, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=2)
    return new_id

def load_chats():
    chats = []
    for fname in sorted(os.listdir(CHAT_DIR), reverse=True):
        if not fname.endswith(".json"):
            continue
        path = os.path.join(CHAT_DIR, fname)
        with open(path, "r", encoding="utf-8") as f:
            msgs = json.load(f)
        chats.append({"id": fname[:-5], "messages": msgs})
    return chats

# ——— Scraper Runner ———
def list_scrapers(scraper_dir="scrapers"):
    return [f[:-3] for f in os.listdir(scraper_dir)
            if f.endswith(".py") and not f.startswith("__")]

common_headers = {
    "sec-ch-ua-platform": "Android",
    "User-Agent": (
        "Mozilla/5.0 (Linux; Android 6.0; Nexus 5 Build/MRA58N)"
        " AppleWebKit/537.36 (KHTML, like Gecko) Chrome/134.0.0.0"
        " Mobile Safari/537.36"
    ),
    "sec-ch-ua": '"Chromium";v="134", "Not:A-Brand";v="24", "Google Chrome";v="134"',
    "sec-ch-ua-mobile": "?1",
}

common_cookies = {
    "_gid": "GA1.3.1046864615.1743240670",
    "monsido": "4081743240673336",
    "_ga_XXXXXXXX": "GS1.1.1743243185.2.1.1743244433.0.0.0",
    "bm_mi": "3C2F88576EB2B1328DF4957982B85D2D~YAAQvvzUF1LjGqOVAQAAWGSd6Bth1EnnrqM/55ra+zZ0CT0o/"
             "5KLuglk/gQSB7kCoLjQwgCbOP906LWlWZpl4fyxcq+yuzGM8msirSFwu1nYdAotFYTHknHGqft33p+"
             "DMIqxmyzvdQzeuYdus7Xtt+oHgGiH8SCgPKX1NtMBWZW5lrG7FfXOfvaS8Odl3AA6lUi25CyUP+fK7"
             "uNQhboYal3H0DmCqbBPi5mqlDApqeGHtAMdQKrVixy2OwbwEhSMMuabDb2ibFZ+tu0ohB4YO1xQHwc"
             "FgoOG6YNswq0nSqtQBryENbhxkjofmazHpE8JywMoO2eWWQm3Txnd52nHkh6EaeI=~1",
    "_gat_gtag_UA_129191797_1": "1",
    "_gat_gtag_UA_67990327_1": "1",
    "_ga_MN793G4JHJ": "GS1.1.1743364571.3.0.1743364571.0.0.0",
    "_ga_FT6SLY9TWT": "GS1.1.1743364376.9.1.1743364574.0.0.0",
    "_ga_0XT7NFV9ZS": "GS1.1.1743364376.9.1.1743364574.0.0.0",
    "bm_sv": "AF7F1D971ACA5FD425CC7DC6D72B9CBC~YAAQvvzUF3PjGqOVAQAAJqGg6Buy7dRTKosyL4YNrqYTl"
             "oJ4Bouxg3EjnJ3fZ0HOiZaZW6nbfsodMC9h0XpffP79Cs0AxpmAR4zH0aL3GIeC4Rhi7ozMlQBhupO"
             "lz+hXJ55VeO7KgaJtW6ym4VjIN/7yh4uk68j3bp+0VK+4ZudN6dkpyRXhfBQXhrNWcT96qjllYRrY"
             "EZ6ZZbPI34HZcdPfFJ0xtuu1BJcV0TFWPeeBL7e3zGyCiwLzvkpECEXA~1",
    "_ga": "GA1.1.1075414505.1743240668",
}

proxies_list = [
    'beqcfgqd:zx2ta8sl24bs@91.217.72.56:6785',
    'beqcfgqd:zx2ta8sl24bs@103.37.181.190:6846',
    'beqcfgqd:zx2ta8sl24bs@45.43.183.159:6471',
    'beqcfgqd:zx2ta8sl24bs@64.137.18.245:6439',
    'beqcfgqd:zx2ta8sl24bs@104.238.50.211:6757',
    'beqcfgqd:zx2ta8sl24bs@89.249.192.133:6532',
    'beqcfgqd:zx2ta8sl24bs@103.101.88.235:5959',
    'beqcfgqd:zx2ta8sl24bs@145.223.45.130:6983',
    'beqcfgqd:zx2ta8sl24bs@45.38.78.112:6049',
]

def fetch_page_with_proxy(
    url,
    proxies_list,
    headers=None,
    cookies=None,
    max_tries=5,
    timeout=10
):
    tried = set()
    attempts = min(max_tries, len(proxies_list))
    for _ in range(attempts):
        proxy = random.choice(proxies_list)
        if proxy in tried:
            continue
        tried.add(proxy)

        proxy_cfg = {
            "http": f"http://{proxy}",
            "https": f"http://{proxy}"
        }

        try:
            resp = requests.get(
                url,
                headers=headers,
                cookies=cookies,
                proxies=proxy_cfg,
                timeout=timeout
            )
            if resp.status_code == 200:
                return resp
            else:
                continue
        except Exception:
            continue

    raise RuntimeError(f"All {len(tried)} proxies failed for {url}")

def get_top_n_listings(
    url,
    proxies,
    headers=None,
    cookies=None,
    n=5,
    selector="a.card",
    max_tries=5,
    timeout=10
):
    resp = fetch_page_with_proxy(
        url, proxies_list=proxies, headers=headers, cookies=cookies,
        max_tries=max_tries, timeout=timeout
    )
    soup = BeautifulSoup(resp.text, "html.parser")
    anchors = soup.select(selector)[:n]
    return [urljoin(url, a["href"]) for a in anchors if a.get("href")]


def run_scraper(module_name, scraper_dir="scrapers"):
    base_name = module_name.split("/")[-1]
    url = FIRST_PAGE_URL.get(base_name)
    meta_json = os.path.join(DATA_DIR, f"{base_name}_top5.json")
    top5 = None

    # Only attempt the “top-5” check for non-RMA sources
    if url and base_name != "RMA":
        selector = TOP5_SELECTORS.get(base_name, DEFAULT_SELECTOR)
        try:
            top5 = get_top_n_listings(
                url,
                proxies=proxies_list,
                headers=common_headers,
                cookies=common_cookies,
                n=5,
                selector=selector
            )
            # If nothing new, skip the scrape
            if os.path.exists(meta_json):
                with open(meta_json, "r") as f:
                    old_top5 = json.load(f)
                if old_top5 == top5:
                    st.sidebar.info("✅ No new updates detected. Skipping scrape.")
                    return
            # Otherwise save the new top-5
            with open(meta_json, "w") as f:
                json.dump(top5, f, indent=2)
        except Exception as e:
            print(f"running full")

    # Run the actual scraper and detect any new CSVs
    before = {f for f in os.listdir() if f.lower().endswith(".csv")}
    script = os.path.join(scraper_dir, f"{module_name}.py")
    try:
        subprocess.run([sys.executable, script], check=True)
    except Exception as e:
        st.sidebar.error(f"Scraper error: {e}")
        return

    after = {f for f in os.listdir() if f.lower().endswith(".csv")}
    new_csvs = after - before

    if new_csvs:
        for csv_file in new_csvs:
            dest = os.path.join(DATA_DIR, csv_file)
            shutil.move(csv_file, dest)
            st.sidebar.success(f"New CSV: {csv_file}")
    else:
        print("no csv produced")


groq_model = "meta-llama/llama-4-scout-17b-16e-instruct"






    # ——— Model Calls ———
async def ask_agent(csv_text: str, question: str, model: str, chat_history: list) -> str:
    import os
    import tiktoken
    from openai import AsyncOpenAI
    from groq import Groq
    import asyncio

    use_groq = model.startswith("meta-llama/")

    try:
        encoding = tiktoken.encoding_for_model(model)
    except KeyError:
        encoding = tiktoken.get_encoding("p50k_base" if use_groq else "cl100k_base")

    def count_tokens(text: str) -> int:
        return len(encoding.encode(text))

    lines = csv_text.strip().split("\n")
    header = lines[0] if lines else ""
    rows = lines[1:] if len(lines) > 1 else []

    MODEL_MAX = 128000 if "gpt-4" in model else 4096
    HEADROOM = 2000
    static_tokens = count_tokens(header) + count_tokens(question)
    usable = MODEL_MAX - HEADROOM - static_tokens

    avg_per_row = max(1, count_tokens("\n".join(rows)) // len(rows)) if rows else 1
    max_rows = max(1, usable // avg_per_row)
    csv_content = "\n".join(rows[:max_rows])

    # ✅ Enhanced system prompt with behavior + examples
    system_prompt = (
    "You are TPI Overwatch AI, a helpful, consistent, and context-aware assistant developed for The Totally and Permanently Incapacitated ex-Service Men and Women’s Association (TPI), a volunteer-led group and the voice for South Australian veterans who have been totally and permanently disabled as a result of their service with the Australian Defence Force. "
    "You write and revise newsletter articles, summaries, and insights for a veteran-focused audience. "
    "You respond in a natural, conversational tone — and only use TPI-style formatting when the user asks for an **article**, **newsletter**, or says 'in the style of TPI'. "
    "In all other cases (like summaries or quick insights), you reply cleanly with minimal formatting."

    "\n\n🧠 Behavior Rules:\n"
    "• If the user prompt contains 'article', 'newsletter', or 'in TPI style', you MUST follow this strict article structure:\n"
    "   – Start with a large `#` heading that includes a fitting emoji and clear title (e.g. `# 🇦🇺 Honouring History and Building Futures`)\n"
    "   – Add a **bold sentence-style subheading** directly beneath the title\n"
    "   – Write a short, italicised intro paragraph that sets the tone and highlights key themes like *honour the past* or *build the future*\n"
    "   – Use `##` subheadings for each major event or topic section (minimum of two is expected)\n"
    "   – Write full narrative-style body text under each section (not bullets)\n"
    "   – Use *italics* and **bold** selectively for rhythm and emphasis\n"
    "   – End the article with a strong closing line — either a reflection or call to action\n"
    "   – Mention SA-specific examples if relevant\n"
    "   – Add helpful inline links if referenced (e.g., [Apply here](https://example.gov.au))\n"
    "   – Total article length must be between 300 and 500 words\n"

    "• If the user prompt does **not** contain 'article', 'newsletter', or 'TPI style', respond in summary format:\n"
    "   – Do **not** use article headings (no `#`, no large title or intro)\n"
    "   – Use **bolded inline section headers** for clarity\n"
    "   – Present updates using clean bullet points or numbered lists\n"
    "   – Be clear, scannable, and professional — avoid long paragraphs or narrative tone\n"

    "\n\n📝 Revision Rules:\n"
    "• Only revise the parts the user requested (e.g. 'make heading bold', 'shorten this')\n"
    "• Return the full updated version (title, subheading, and body)\n"
    "• Respect tone/style requests (e.g. 'friendlier', 'more concise')\n"
    "• Apply multiple edits if asked\n"
    "• Ask for clarification if instructions are vague"

    "\n\n📋 Examples (FORMAT ONLY — DO NOT COPY CONTENT):\n"

    "\nUser: Write a TPI newsletter article about a new support program\n"
    "Assistant:\n"
    "# 🛡 Expanding Support: $4.8M Boost for Veteran Advocacy\n"
    "**New funding announced for ex-service organisations across Australia**\n"
    "*In a powerful move to strengthen veteran support services, the Department of Veterans' Affairs has launched...*\n\n"
    "## What’s Available?\n"
    "The government has announced $4.8 million in funding through the BEST Grants Program to help veteran organisations...\n\n"
    "## How to Apply\n"
    "Applications close on 28 April 2025. Grants cover legal support, advocacy training, and assistance with claims...\n\n"
    "*Let’s honour their legacy by amplifying their voice.*"

    "\nUser: Summarize key veteran updates\n"
    "Assistant:\n"
    "**Honouring Our Korean War Veterans** – Marks 75 years since Australia’s involvement.\n"
    "**Employment Awards** – Nominations open for 2025.\n"
    "**Grants Program** – $3.5M in community funding available.\n"
    "**Financial Boost** – Over 150,000 veterans to receive increased compensation.\n"
    "**Advocacy Support** – $4.8M in BEST grants for ex-service organisations.\n"

    "\n\nCurrent date: July 2025\nKnowledge cutoff: June 2024"
)

    # 👤 Final user message: includes actual CSV
    user_prompt = f"""
Here is the dataset and question.

CSV:
{header}
{csv_content}

Question: {question}
"""

    # 🔁 Full chat message array
    messages = [{"role": "system", "content": system_prompt}]
    if chat_history:
        for item in chat_history[-40:]:
            if isinstance(item, dict) and 'role' in item and 'content' in item:
                messages.append({"role": item['role'], "content": item['content']})
    messages.append({"role": "user", "content": user_prompt})


    

    # ——— Model Calls ———
    async def send_openai() -> str:
        client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

        # Add tool if using search-preview model
        use_search_tool = "search-preview" in model
        extra_args = {"tools": [{"type": "web_search"}]} if use_search_tool else {}

        resp = await client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.7 if "article" in question.lower() else 0.1,
            top_p=0.95 if "article" in question.lower() else 0.1,
            max_tokens=4000,
            **extra_args  # ⬅️ Web search tool injected here
        )
        return resp.choices[0].message.content.strip()


    async def send_groq() -> str:
        def sync():
            client = Groq(api_key=os.getenv("GROQ_API_KEY"))
            resp = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=0.7 if "article" in question.lower() else 0.1,
                top_p=0.95 if "article" in question.lower() else 0.1,
                max_tokens=4000
            )
            return resp.choices[0].message.content.strip()
        return await asyncio.to_thread(sync)

    return await send_groq() if use_groq else await send_openai()


    
def create_docx(content: str) -> BytesIO:
    doc = Document()
    doc.add_heading("Assistant Response", level=1)
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def main():

    st.set_page_config(page_title="TPI Overwatch")
    
    st.markdown("""
    <style>
      .user-message {
        border: 1px solid #6B7280;
        color: #4B5563;
        padding: 0.75rem 1rem;
        border-radius: 0.375rem;
        margin: 0.5rem 0;
        max-width: 80%;
      }
      .assistant-message {
        color: #F9FAFB;
        padding: 0.75rem 1rem;
        border-radius: 0.375rem;
        margin: 0.5rem 0;
        max-width: 80%;
      }
      .assistant-message h1,
      .assistant-message h2,
      .assistant-message h3,
      .assistant-message h4,
      .assistant-message h5,
      .assistant-message h6 {
        color: #FFFFFF !important;
        margin-top: 0.8rem;
        margin-bottom: 0.4rem;
      }
      .assistant-message ul {
        margin-left: 1.5rem;
      }
    </style>
    """, unsafe_allow_html=True)



    st.title("🕸️ TPI Overwatch")

    if "query" not in st.session_state:
        st.session_state["query"] = ""

    st.sidebar.image("logo.png", width=200)

    source_groups = {
            "Department of Veteran Affairs (DVA)": [
                "DVA Minister",
                "DVA Veteran Affairs",
                "DVA Repatriation Commission",
                "DVA Website About",
                "DVA Website Home",
                "DVA Website Latest News",
                # "X DVA",
                # "Instagram DVA"
            ],
            "Australian War Memorial (AWM)": [
                "articles",
                # "X AWM",
                # "Instagram AWM"
            ],
            "Repatriation Medical Authority (RMA)": [
                "RMA",
            ],
        }


    choice = st.sidebar.selectbox("Select Source", list(source_groups.keys()))

    if st.sidebar.button("Fetch Latest Content"):
        before = set(os.listdir(DATA_DIR))
        for name in source_groups[choice]:
            module_path = f"{choice}/{name}" 
            run_scraper(module_path)
        # move any new CSVs into data/
        for f in set(os.listdir()) - before:
            if f.endswith(".csv"):
                shutil.move(f, os.path.join(DATA_DIR, f))
                st.sidebar.success(f"New CSV: {f}")

    st.sidebar.markdown("---")

    # 2) Now pick *which* group you want to load data from
    focus_choice = st.sidebar.selectbox("Focus", list(source_groups.keys()))

    # 2) list actual .csv files in data/ matching that group
    # list all the .csv files in data/
    all_csvs = [f for f in os.listdir(DATA_DIR) if f.endswith(".csv")]
    # filter to just those belonging to the current focus group
    matching = [f for f in all_csvs if any(f.startswith(name) for name in source_groups[focus_choice])]

    if not matching:
        st.sidebar.warning(f"No CSVs found for {focus_choice}.")
        return

    # strip off “.csv” for display
    options = [os.path.splitext(f)[0] for f in matching]

    # single‐select radio
    sel_file = st.sidebar.radio(f"Select one {focus_choice} CSV", options)

    # load the chosen file (add the extension back)
    if "last_csv_selected" not in st.session_state:
        st.session_state["last_csv_selected"] = None


    csv_path = os.path.join(DATA_DIR, f"{sel_file}.csv")


    if sel_file != st.session_state["last_csv_selected"]:
        st.session_state["chat_id"] = None
        st.session_state["chat_history"] = []
        st.session_state["new_chat_title"] = None
        st.session_state["last_csv_selected"] = sel_file

    if os.path.exists(csv_path):
        try:
            df = pd.read_csv(csv_path)
        except pd.errors.EmptyDataError:
            st.error("❌ The CSV file exists but is empty or has no columns. Scraping may have failed or been blocked.")
            st.stop()
    else:
        st.error("❌ CSV file not found.")
        st.stop()

    chats = load_chats()
    options = ["🆕 New Article Thread"] + [
        c.get("title", c['id']) for c in chats
    ]
    default_idx = 0
    if st.session_state.get("chat_id"):
        default_idx = next(
            (i+1 for i, c in enumerate(chats) if c["id"] == st.session_state.chat_id),
            0
        )
    sel = st.sidebar.selectbox(
        "Article Threads",
        options,
        index=default_idx,
        key="chat_select"
    )
    if sel == "🆕 New Article Thread":
        st.session_state.chat_id = None
        st.session_state.chat_history = []
        st.session_state["new_chat_title"] = None
    else:
        idx = options.index(sel) - 1
        st.session_state.chat_id = chats[idx]["id"]
        st.session_state.chat_history = chats[idx]["messages"].copy()

    st.subheader(f"Dataset: {sel_file}")
    st.dataframe(df)
    st.markdown("---")

    st.sidebar.header("Predefined Prompts")
    predef_prompts = [
        "Summarize the key insights from this dataset.",
        "Write a comprehensive article based on this dataset, weaving in key insights, context, and potential implications.",
        "Give me a narrative overview of what this data represents."
    ]
    def _set_query(p):
        st.session_state["query"] = p

    for i, p in enumerate(predef_prompts):
        st.sidebar.button(
            p,
            key=f"predef_{i}",
            on_click=_set_query,
            args=(p,)
        )

    # Model selectbox above text input box
    raw_model = st.selectbox(
        "Model",
        ["gpt-4o", "Groq"],
        key="model_select"
    )
    model = raw_model if raw_model != "Groq" else "meta-llama/llama-4-scout-17b-16e-instruct"

    with st.form("chat_form", clear_on_submit=True):
        query = st.text_input("Ask anything—article, summary, insight…", key="query")


        submitted = st.form_submit_button("Ask Agent")
        if submitted and query:
            st.session_state.chat_history.append({"role": "user", "content": query})
            csv_text = df.to_csv(index=False)
            with st.spinner("🤖 Agent is thinking..."):
                answer = asyncio.run(ask_agent(csv_text, query, model, st.session_state.chat_history))
            st.session_state.chat_history.append({"role": "assistant", "content": answer})
            csv_basename = os.path.splitext(sel_file)[0]
            if st.session_state.get("new_chat_title") is None:
                st.session_state["new_chat_title"] = query[:50]

            chat_title = st.session_state["new_chat_title"]

            new_id = save_chat(
                st.session_state.chat_history,
                chat_id=st.session_state.get("chat_id"),
                title=chat_title
            )
            st.session_state.chat_id = new_id
            
            try:
                st.experimental_rerun()
            except AttributeError:
                st.rerun()

    # ——— Display Chat in Descending Order by Pairs ———
    chat_history = st.session_state.chat_history
    pairs = []
    for i in range(0, len(chat_history), 2):
        if i + 1 < len(chat_history):  
            pairs.append((chat_history[i], chat_history[i + 1]))
        else: 
            pairs.append((chat_history[i], None))
    for pair_idx, (user_msg, assistant_msg) in enumerate(reversed(pairs)):

        st.markdown(f'<div class="user-message">👤 {user_msg["content"]}</div>', unsafe_allow_html=True)
    
        if assistant_msg and assistant_msg["content"].strip():
            with st.container():
                st.markdown('<div class="assistant-message">', unsafe_allow_html=True)
                st.markdown(assistant_msg["content"], unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)
        
                docx_buffer = create_docx(assistant_msg["content"])
                st.download_button(
                    label="Save Article Draft",
                    data=docx_buffer,
                    file_name=f"Article_{pair_idx}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{pair_idx}"
                )

    st.markdown("---")





if __name__ == "__main__":
    main()
